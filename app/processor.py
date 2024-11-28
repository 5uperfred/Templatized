from docx import Document
from pdf2docx import Converter
import PyPDF2
import json
import os
import openai
from typing import Dict, List, Any, Tuple
import logging
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import tempfile
import magic
from tqdm import tqdm

class DocumentProcessor:
    def __init__(self, api_key: str):
        """Initialize the processor with OpenAI API key."""
        self.api_key = api_key
        openai.api_key = api_key
        self.logger = logging.getLogger(__name__)

    def detect_file_type(self, file_path: str) -> str:
        """Detect the type of file using magic numbers."""
        mime = magic.Magic(mime=True)
        file_type = mime.from_file(file_path)
        return file_type

    def pdf_to_docx(self, pdf_path: str) -> str:
        """Convert PDF to DOCX format."""
        try:
            # Create temporary file for DOCX
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                docx_path = tmp_docx.name

            # Convert PDF to DOCX
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()

            return docx_path
        except Exception as e:
            self.logger.error(f"Error converting PDF to DOCX: {str(e)}")
            raise

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF for variable identification."""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            self.logger.error(f"Error extracting text from PDF: {str(e)}")
            raise

    def capture_run_formatting(self, run) -> Dict[str, Any]:
        """Capture all formatting details of a run."""
        return {
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size.pt if run.font.size else None,
            'font_color': run.font.color.rgb if run.font.color else None,
            'highlight_color': run.font.highlight_color if hasattr(run.font, 'highlight_color') else None,
            'strike': run.font.strike,
            'subscript': run.font.subscript,
            'superscript': run.font.superscript
        }

    def capture_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """Capture all formatting details of a paragraph."""
        return {
            'alignment': paragraph.alignment,
            'style': paragraph.style.name,
            'line_spacing': paragraph.paragraph_format.line_spacing,
            'space_before': paragraph.paragraph_format.space_before.pt if paragraph.paragraph_format.space_before else None,
            'space_after': paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else None,
            'first_line_indent': paragraph.paragraph_format.first_line_indent.pt if paragraph.paragraph_format.first_line_indent else None,
            'left_indent': paragraph.paragraph_format.left_indent.pt if paragraph.paragraph_format.left_indent else None,
            'right_indent': paragraph.paragraph_format.right_indent.pt if paragraph.paragraph_format.right_indent else None,
            'runs': [self.capture_run_formatting(run) for run in paragraph.runs]
        }

    def extract_variables(self, document_text: str) -> Dict[str, str]:
        """Extract variables using OpenAI's GPT model."""
        try:
            prompt = """
            Analyze this document and identify all variables.
            Return a Python dictionary where:
            - Keys are the original text
            - Values are standardized variable names
            Focus on:
            - Names and titles
            - Dates
            - Monetary amounts
            - Percentages
            - Account numbers
            - Addresses
            - Legal entity names
            - Contact information
            - Reference numbers
            - Terms and conditions variables
            Return only the Python dictionary, no additional text.
            """

            response = openai.ChatCompletion.create(
                model="gpt-4o-mini-2024-07-18",
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": document_text}
                ],
                temperature=0.1
            )

            variables_dict = eval(response.choices[0].message.content)
            return variables_dict
        except Exception as e:
            self.logger.error(f"Error extracting variables: {str(e)}")
            raise

    def create_template_document(self, template_text: List[Dict[str, Any]], output_path: str):
        """Create a new Word document with all formatting preserved."""
        doc = Document()

        for para_info in template_text:
            paragraph = doc.add_paragraph()
            paragraph.style = para_info['style']
            paragraph.alignment = para_info['alignment']

            pf = paragraph.paragraph_format
            if para_info['line_spacing']:
                pf.line_spacing = para_info['line_spacing']
            if para_info['space_before']:
                pf.space_before = Pt(para_info['space_before'])
            if para_info['space_after']:
                pf.space_after = Pt(para_info['space_after'])
            if para_info['first_line_indent']:
                pf.first_line_indent = Pt(para_info['first_line_indent'])
            if para_info['left_indent']:
                pf.left_indent = Pt(para_info['left_indent'])
            if para_info['right_indent']:
                pf.right_indent = Pt(para_info['right_indent'])

            for run_info in para_info['runs']:
                run = paragraph.add_run(run_info['text'])
                run.bold = run_info['bold']
                run.italic = run_info['italic']
                run.underline = run_info['underline']
                run.font.name = run_info['font_name']
                if run_info['font_size']:
                    run.font.size = Pt(run_info['font_size'])
                if run_info['font_color']:
                    run.font.color.rgb = run_info['font_color']
                run.font.strike = run_info['strike']
                run.font.subscript = run_info['subscript']
                run.font.superscript = run_info['superscript']

        doc.save(output_path)

    def process_document(self, file_path: str, output_dir: str) -> Tuple[str, str]:
        """Process either PDF or DOCX document and generate template."""
        try:
            os.makedirs(output_dir, exist_ok=True)

            # Detect file type
            file_type = self.detect_file_type(file_path)

            # Convert PDF to DOCX if necessary
            if 'pdf' in file_type.lower():
                self.logger.info("Converting PDF to DOCX...")
                docx_path = self.pdf_to_docx(file_path)
                original_text = self.extract_text_from_pdf(file_path)
            else:
                docx_path = file_path
                doc = Document(docx_path)
                original_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

            # Process document
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            template_path = os.path.join(output_dir, f"{base_name}_template.docx")
            variables_path = os.path.join(output_dir, f"{base_name}_variables.json")

            # Extract variables
            self.logger.info("Extracting variables...")
            variables_dict = self.extract_variables(original_text)

            # Create template
            self.logger.info("Creating template...")
            doc = Document(docx_path)
            formatted_text = [self.capture_paragraph_formatting(p) for p in doc.paragraphs]

            # Apply variables while preserving formatting
            for paragraph in formatted_text:
                for run in paragraph['runs']:
                    for original, variable_name in variables_dict.items():
                        if original in run['text']:
                            run['text'] = run['text'].replace(original, f"{{{{ {variable_name} }}}}")

            # Create final documents
            self.create_template_document(formatted_text, template_path)

            with open(variables_path, 'w', encoding='utf-8') as f:
                json.dump(variables_dict, f, indent=4)

            # Cleanup temporary files
            if 'pdf' in file_type.lower() and docx_path != file_path:
                os.remove(docx_path)

            return template_path, variables_path

        except Exception as e:
            self.logger.error(f"Error processing document: {str(e)}")
            raise
